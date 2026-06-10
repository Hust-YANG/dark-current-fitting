#!/usr/bin/env python3
"""
Dark Current Component Fitting Script
======================================

Dual-model support:
  Model 1 (Eq.1): J_dark = J0*[exp(qV/(A*kT))-1] + V/Rsh + k*V^m
  Model 2 (Eq.2): J_dark = J0*[exp(qV/(A*kT))-1] + V/Rsh + B*V*exp(-c/(Vbi-V))

Voltage convention: raw data V>0 reverse bias, V<0 forward bias.
Script negates: V_fit = -V_raw, J_fit = -J_raw.

Output:
  1. Fitting plots (SVG/PDF/PNG@300dpi/PNG@600dpi)
  2. Component data .txt files
  3. Fitting parameters CSV
  4. Combined academic report .docx
"""

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.optimize import curve_fit
from scipy.constants import k, e
import pandas as pd
import os, sys, json, io
import warnings
warnings.filterwarnings('ignore', category=UserWarning, module='matplotlib')
from datetime import datetime

if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

T_DEFAULT = 300
CM_PER_INCH = 2.54

COLOR_PALETTE = {
    'data':       '#0072B2',
    'total_fit':  '#D55E00',
    'j_main':     '#009E73',
    'j_ohm':      '#CC79A7',
    'j_non_tat':  '#56B4E9',
}

BOUNDS_BASE = (
    [1e-15, 1.0,  1e-1,  1e-12, 0.001, 0.01],
    [1e-2,  3.0,  1e12,  1e8,   100.0,  2.0],
)
P0_DEFAULT    = [1e-7,  1.5,  1e4,  1.0,   1.0,  0.5]
P0_ALT        = [1e-5,  1.2,  1e5,  10.0,  5.0,  0.8]
P0_AGGRESSIVE = [1e-6,  1.8,  1e3,  0.1,   0.5,  0.4]

BOUNDS_M1 = (
    [1e-15, 1.0,  1e-1,  1e-15, 0.5],
    [1e-2,  3.0,  1e12,  1e2,   5.0],
)
P0_DEFAULT_M1    = [1e-7,  1.5,  1e4,  1e-8,  2.0]
P0_ALT_M1        = [1e-5,  1.2,  1e5,  1e-6,  2.5]
P0_AGGRESSIVE_M1 = [1e-6,  1.8,  1e3,  1e-10, 1.5]

V_SEG_DEFAULT = 0.2

CONSISTENCY_RULES = {
    'J0':  'decrease', 'A':   'decrease', 'Rsh': 'increase',
    'B':   'decrease', 'c':   'increase',
}
CONSISTENCY_RULES_M1 = {
    'J0':  'decrease', 'A':   'decrease', 'Rsh': 'increase',
    'k':   'decrease',
}

def _load_reference_data():
    ref_dir = os.path.dirname(os.path.abspath(__file__))
    ref_dir = os.path.join(os.path.dirname(ref_dir), 'references')
    ref_path = os.path.join(ref_dir, 'reference_parameters.json')
    if os.path.exists(ref_path):
        try:
            with open(ref_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return None

REFERENCE_DATA = _load_reference_data()

def thermal_voltage(T):
    return k * T / e

def main_diode_current(V, J0, A, Vt):
    return J0 * (np.exp(V / (A * Vt)) - 1.0)

def ohmic_current(V, Rsh):
    return V / Rsh

def tat_current(V, B, c, Vbi):
    denom = Vbi - V
    valid = denom > 0
    result = np.zeros_like(V, dtype=np.float64)
    if np.any(valid):
        exponent = np.clip(-c / denom[valid], -700.0, 0.0)
        result[valid] = B * V[valid] * np.exp(exponent)
    return result

def three_component_model(V, J0, A, Rsh, B, c, Vbi, Vt=None, T=300):
    if Vt is None:
        Vt = thermal_voltage(T)
    return (main_diode_current(V, J0, A, Vt) + ohmic_current(V, Rsh) + tat_current(V, B, c, Vbi))

def non_ohmic_tunneling(V, k, m):
    return k * np.sign(V) * np.power(np.abs(V), m)

def model1(V, J0, A, Rsh, k, m, Vt=None, T=300):
    if Vt is None:
        Vt = thermal_voltage(T)
    return (main_diode_current(V, J0, A, Vt) + ohmic_current(V, Rsh) + non_ohmic_tunneling(V, k, m))

def detect_encoding(filepath):
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'gb18030', 'utf-16', 'utf-16-le', 'utf-16-be', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                f.read(1024)
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return 'utf-8'

def _parse_raw_data(filepath):
    if not os.path.exists(filepath):
        print(f"Error: file not found {filepath}")
        sys.exit(1)
    encoding = detect_encoding(filepath)
    sweeps = {}
    flat_V, flat_I = [], []
    has_index = False
    with open(filepath, 'r', encoding=encoding) as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            parts = [p.strip().strip('"') for p in line.replace('\t', ' ').split(' ') if p.strip()]
            if len(parts) < 2:
                continue
            try:
                if len(parts) >= 3:
                    idx_str = parts[0]
                    V_pt, I_pt = float(parts[-2]), float(parts[-1])
                    if '/' in idx_str:
                        sweep_id = idx_str.split('/')[0]
                        has_index = True
                    else:
                        sweep_id = '1'
                    if sweep_id not in sweeps:
                        sweeps[sweep_id] = {'V': [], 'I': []}
                    sweeps[sweep_id]['V'].append(V_pt)
                    sweeps[sweep_id]['I'].append(I_pt)
                else:
                    V_pt, I_pt = float(parts[0]), float(parts[1])
                    flat_V.append(V_pt)
                    flat_I.append(I_pt)
            except ValueError:
                continue
    if has_index and len(sweeps) > 0:
        return sweeps, True
    elif len(flat_V) > 0:
        return {'1': {'V': flat_V, 'I': flat_I}}, False
    elif len(sweeps) > 0:
        return sweeps, False
    else:
        print(f"Error: no valid data in {filepath}")
        sys.exit(1)

def _identify_dark_sweeps(sweeps_dict, filepath):
    dark_sweeps = []
    light_sweeps = []
    print(f"\n  Sweep detection ({os.path.basename(filepath)}):")
    print(f"  {'Sweep':<8} {'Pts':<6} {'I@V0':<16} {'max|I|':<16} {'Ratio':<10} {'State'}")
    for sweep_id in sorted(sweeps_dict.keys()):
        V_arr = np.array(sweeps_dict[sweep_id]['V'])
        I_arr = np.array(sweeps_dict[sweep_id]['I'])
        idx_zero = np.argmin(np.abs(V_arr))
        I_at_zero = I_arr[idx_zero]
        max_abs_I = np.max(np.abs(I_arr))
        if max_abs_I > 0:
            ratio = abs(I_at_zero) / max_abs_I
        else:
            ratio = 0.0
        is_dark = (ratio < 0.01) or (abs(I_at_zero) < 1e-12)
        state = 'DARK' if is_dark else 'LIGHT (skip)'
        print(f"  {sweep_id:<8} {len(V_arr):<6} {I_at_zero:<16.4e} {max_abs_I:<16.4e} {ratio:<10.4f} {state}")
        if is_dark:
            dark_sweeps.append(sweep_id)
        else:
            light_sweeps.append(sweep_id)
    if not dark_sweeps:
        print(f"  Warning: No dark sweep detected, using all data")
        dark_sweeps = list(sweeps_dict.keys())
    print(f"  -> Using dark sweep(s): {dark_sweeps}")
    if light_sweeps:
        print(f"  -> Discarded light sweep(s): {light_sweeps}")
    return dark_sweeps

def load_data(filepath, area=None, auto_dark=True, manual_sweep=None, max_points=None):
    sweeps_dict, has_index = _parse_raw_data(filepath)
    if manual_sweep is not None:
        if manual_sweep in sweeps_dict:
            use_sweeps = [manual_sweep]
        else:
            use_sweeps = _identify_dark_sweeps(sweeps_dict, filepath) if auto_dark else list(sweeps_dict.keys())
    elif auto_dark and len(sweeps_dict) > 1:
        use_sweeps = _identify_dark_sweeps(sweeps_dict, filepath)
    else:
        use_sweeps = list(sweeps_dict.keys())
    all_V, all_I = [], []
    for sid in sorted(use_sweeps):
        all_V.extend(sweeps_dict[sid]['V'])
        all_I.extend(sweeps_dict[sid]['I'])
    if max_points is not None and max_points > 0:
        all_V = all_V[:max_points]
        all_I = all_I[:max_points]
    V_raw = np.array(all_V)
    I_raw = np.array(all_I)
    V_fit = -V_raw
    if area is not None and area > 0:
        J_fit = -I_raw / area
    else:
        J_fit = -I_raw
    print(f"  Loaded {len(V_fit)} pts")
    return V_fit, J_fit

def fit_dark_current(V, J, sample_name="Sample", T=300):
    Vt = thermal_voltage(T)
    print(f"\nFitting (Eq.2): {sample_name}")
    def fit_func(V, J0, A, Rsh, B, c, Vbi):
        return three_component_model(V, J0, A, Rsh, B, c, Vbi, Vt=Vt)
    for i, p0 in enumerate([P0_DEFAULT, P0_ALT, P0_AGGRESSIVE]):
        try:
            popt, pcov = curve_fit(fit_func, V, J, p0=p0, bounds=BOUNDS_BASE, maxfev=100000, method='trf')
            J_pred = fit_func(V, *popt)
            ss_res = np.sum((J - J_pred) ** 2)
            ss_tot = np.sum((J - np.mean(J)) ** 2)
            r_squared = 1 - ss_res / ss_tot
            perr = np.sqrt(np.diag(pcov))
            print(f"  J0={popt[0]:.4e} A={popt[1]:.4f} Rsh={popt[2]:.2e} B={popt[3]:.4f} c={popt[4]:.4f} Vbi={popt[5]:.4f} R2={r_squared:.6f}")
            return {'popt': popt, 'pcov': pcov, 'perr': perr, 'r_squared': r_squared, 'Vt': Vt, 'model': 'eq48'}
        except RuntimeError:
            pass
    print("All fits failed.")
    return None

def segmented_fit_dark_current(V, J, sample_name="Sample", T=300, V_seg=V_SEG_DEFAULT, verbose=True):
    Vt = thermal_voltage(T)
    global_seed = fit_dark_current(V, J, sample_name, T)
    if global_seed is not None:
        gJ0, gA, gRsh, gB, gc, gVbi = global_seed['popt']
    else:
        gJ0, gA, gRsh, gB, gc, gVbi = P0_DEFAULT
    mask_small = np.abs(V) < V_seg
    V_small, J_small = V[mask_small], J[mask_small]
    def stage1_model(V, J0, A, Rsh_eff):
        return main_diode_current(V, J0, A, Vt) + ohmic_current(V, Rsh_eff)
    p0_options_s1 = [[gJ0, gA, gRsh], [P0_DEFAULT[0], P0_DEFAULT[1], P0_DEFAULT[2]],
                     [P0_ALT[0], P0_ALT[1], P0_ALT[2]], [P0_AGGRESSIVE[0], P0_AGGRESSIVE[1], P0_AGGRESSIVE[2]]]
    bounds_s1 = ([BOUNDS_BASE[0][0], BOUNDS_BASE[0][1], BOUNDS_BASE[0][2]],
                 [BOUNDS_BASE[1][0], BOUNDS_BASE[1][1], BOUNDS_BASE[1][2]])
    s1_success = False; r2_s1 = None
    for guess_idx, p0_guess in enumerate(p0_options_s1):
        try:
            popt_s1, pcov_s1 = curve_fit(stage1_model, V_small, J_small, p0=p0_guess, bounds=bounds_s1, maxfev=50000, method='trf')
            J_pred_s1 = stage1_model(V_small, *popt_s1)
            ss_res_s1 = np.sum((J_small - J_pred_s1) ** 2)
            ss_tot_s1 = np.sum((J_small - np.mean(J_small)) ** 2)
            r2_s1 = 1 - ss_res_s1 / ss_tot_s1 if ss_tot_s1 > 0 else 0.0
            s1_success = True; break
        except RuntimeError:
            pass
    if not s1_success:
        global_fit = fit_dark_current(V, J, sample_name, T)
        if global_fit is None: return None
        J0_s1, A_s1 = global_fit['popt'][0], global_fit['popt'][1]; pcov_s1 = None
    else:
        J0_s1, A_s1, Rsh_eff = popt_s1
    mask_large = V < -V_seg
    V_large, J_large = V[mask_large], J[mask_large]
    if len(V_large) < 5:
        mask_large = V < 0; V_large, J_large = V[mask_large], J[mask_large]
    def stage2_model(V, Rsh, B, c, Vbi):
        return (main_diode_current(V, J0_s1, A_s1, Vt) + ohmic_current(V, Rsh) + tat_current(V, B, c, Vbi))
    p0_options_s2 = [[gRsh, gB, gc, gVbi], [P0_DEFAULT[2], P0_DEFAULT[3], P0_DEFAULT[4], P0_DEFAULT[5]],
                     [P0_ALT[2], P0_ALT[3], P0_ALT[4], P0_ALT[5]], [P0_AGGRESSIVE[2], P0_AGGRESSIVE[3], P0_AGGRESSIVE[4], P0_AGGRESSIVE[5]]]
    bounds_s2 = ([BOUNDS_BASE[0][2], BOUNDS_BASE[0][3], BOUNDS_BASE[0][4], BOUNDS_BASE[0][5]],
                 [BOUNDS_BASE[1][2], BOUNDS_BASE[1][3], BOUNDS_BASE[1][4], BOUNDS_BASE[1][5]])
    s2_success = False; r2_s2 = None
    for guess_idx, p0_guess in enumerate(p0_options_s2):
        try:
            popt_s2, pcov_s2 = curve_fit(stage2_model, V_large, J_large, p0=p0_guess, bounds=bounds_s2, maxfev=100000, method='trf')
            J_pred_s2 = stage2_model(V_large, *popt_s2)
            ss_res_s2 = np.sum((J_large - J_pred_s2) ** 2)
            ss_tot_s2 = np.sum((J_large - np.mean(J_large)) ** 2)
            r2_s2 = 1 - ss_res_s2 / ss_tot_s2 if ss_tot_s2 > 0 else 0.0
            s2_success = True; break
        except RuntimeError:
            pass
    if not s2_success:
        global_fit = fit_dark_current(V, J, sample_name, T)
        if global_fit is None: return None
        Rsh_s2 = global_fit['popt'][2]; B_s2 = global_fit['popt'][3]; c_s2 = global_fit['popt'][4]; Vbi_s2 = global_fit['popt'][5]; pcov_s2 = None
    else:
        Rsh_s2, B_s2, c_s2, Vbi_s2 = popt_s2
    p0_s3 = [J0_s1, A_s1, Rsh_s2, B_s2, c_s2, Vbi_s2]
    bounds_s3 = (
        [max(BOUNDS_BASE[0][0], J0_s1 * 0.1), max(BOUNDS_BASE[0][1], A_s1 * 0.7),
         max(BOUNDS_BASE[0][2], Rsh_s2 * 0.1), max(BOUNDS_BASE[0][3], B_s2 * 0.01),
         max(BOUNDS_BASE[0][4], c_s2 * 0.1), max(BOUNDS_BASE[0][5], Vbi_s2 * 0.5)],
        [min(BOUNDS_BASE[1][0], J0_s1 * 10.0), min(BOUNDS_BASE[1][1], A_s1 * 1.5),
         min(BOUNDS_BASE[1][2], Rsh_s2 * 10.0), min(BOUNDS_BASE[1][3], B_s2 * 100.0),
         min(BOUNDS_BASE[1][4], c_s2 * 10.0), min(BOUNDS_BASE[1][5], Vbi_s2 * 1.5)],
    )
    s3_success = False
    try:
        popt, pcov = curve_fit(lambda V, *p: three_component_model(V, *p, Vt=Vt), V, J, p0=p0_s3, bounds=bounds_s3, maxfev=50000, method='trf')
        perr = np.sqrt(np.diag(pcov))
        s3_success = True
    except RuntimeError:
        pass
    if not s3_success:
        popt = np.array(p0_s3)
        perr = np.full(6, np.nan) if not (s1_success and s2_success) else np.array([
            np.sqrt(pcov_s1[0, 0]), np.sqrt(pcov_s1[1, 1]), np.sqrt(pcov_s2[0, 0]),
            np.sqrt(pcov_s2[1, 1]), np.sqrt(pcov_s2[2, 2]), np.sqrt(pcov_s2[3, 3])])
    J_pred = three_component_model(V, *popt, Vt=Vt)
    ss_res = np.sum((J - J_pred) ** 2)
    ss_tot = np.sum((J - np.mean(J)) ** 2)
    r_squared = 1 - ss_res / ss_tot if ss_tot > 0 else 0.0
    return {'popt': popt, 'pcov': pcov if s3_success else None, 'perr': perr,
            'r_squared': r_squared, 'Vt': Vt, 'model': 'eq48',
            'segmented': True, 'V_seg': V_seg,
            'stage1': {'J0': J0_s1, 'A': A_s1, 'r2': r2_s1},
            'stage2': {'Rsh': Rsh_s2, 'B': B_s2, 'c': c_s2, 'Vbi': Vbi_s2, 'r2': r2_s2}}

def fit_model1(V, J, sample_name="Sample", T=300):
    Vt = thermal_voltage(T)
    print(f"\nFitting (Eq.1): {sample_name}")
    def fit_func(V, J0, A, Rsh, k, m):
        return model1(V, J0, A, Rsh, k, m, Vt=Vt)
    for i, p0 in enumerate([P0_DEFAULT_M1, P0_ALT_M1, P0_AGGRESSIVE_M1]):
        try:
            popt, pcov = curve_fit(fit_func, V, J, p0=p0, bounds=BOUNDS_M1, maxfev=100000, method='trf')
            J_pred = fit_func(V, *popt)
            ss_res = np.sum((J - J_pred) ** 2)
            ss_tot = np.sum((J - np.mean(J)) ** 2)
            r_squared = 1 - ss_res / ss_tot
            perr = np.sqrt(np.diag(pcov))
            print(f"  J0={popt[0]:.4e} A={popt[1]:.4f} Rsh={popt[2]:.2e} k={popt[3]:.4e} m={popt[4]:.4f} R2={r_squared:.6f}")
            return {'popt': popt, 'pcov': pcov, 'perr': perr, 'r_squared': r_squared, 'Vt': Vt, 'model': 'eq1'}
        except RuntimeError:
            pass
    print("All fits failed.")
    return None

def segmented_fit_model1(V, J, sample_name="Sample", T=300, V_seg=V_SEG_DEFAULT, verbose=True):
    Vt = thermal_voltage(T)
    global_seed = fit_model1(V, J, sample_name, T)
    if global_seed is not None:
        gJ0, gA, gRsh, gk, gm = global_seed['popt']
    else:
        gJ0, gA, gRsh, gk, gm = P0_DEFAULT_M1
    mask_small = np.abs(V) < V_seg
    V_small, J_small = V[mask_small], J[mask_small]
    def stage1_model(V, J0, A, Rsh_eff):
        return main_diode_current(V, J0, A, Vt) + ohmic_current(V, Rsh_eff)
    p0_opts = [[gJ0, gA, gRsh], [P0_DEFAULT_M1[0], P0_DEFAULT_M1[1], P0_DEFAULT_M1[2]],
               [P0_ALT_M1[0], P0_ALT_M1[1], P0_ALT_M1[2]], [P0_AGGRESSIVE_M1[0], P0_AGGRESSIVE_M1[1], P0_AGGRESSIVE_M1[2]]]
    bnd = ([BOUNDS_M1[0][0], BOUNDS_M1[0][1], BOUNDS_M1[0][2]], [BOUNDS_M1[1][0], BOUNDS_M1[1][1], BOUNDS_M1[1][2]])
    s1_ok = False; r2_s1 = None
    for gi, pg in enumerate(p0_opts):
        try:
            popt1, pcov1 = curve_fit(stage1_model, V_small, J_small, p0=pg, bounds=bnd, maxfev=50000, method='trf')
            Jp1 = stage1_model(V_small, *popt1)
            ssr1 = np.sum((J_small - Jp1)**2); sst1 = np.sum((J_small - np.mean(J_small))**2)
            r2_s1 = 1 - ssr1 / sst1 if sst1 > 0 else 0.0
            s1_ok = True; break
        except RuntimeError:
            pass
    if not s1_ok:
        gf = fit_model1(V, J, sample_name, T)
        if gf is None: return None
        J0_s1, A_s1 = gf['popt'][0], gf['popt'][1]
    else:
        J0_s1, A_s1, _ = popt1
    mask_large = V < -V_seg
    V_large, J_large = V[mask_large], J[mask_large]
    if len(V_large) < 5:
        mask_large = V < 0; V_large, J_large = V[mask_large], J[mask_large]
    def stage2_model(V, Rsh, k, m):
        return (main_diode_current(V, J0_s1, A_s1, Vt) + ohmic_current(V, Rsh) + non_ohmic_tunneling(V, k, m))
    p0_opts2 = [[gRsh, gk, gm], [P0_DEFAULT_M1[2], P0_DEFAULT_M1[3], P0_DEFAULT_M1[4]],
                [P0_ALT_M1[2], P0_ALT_M1[3], P0_ALT_M1[4]], [P0_AGGRESSIVE_M1[2], P0_AGGRESSIVE_M1[3], P0_AGGRESSIVE_M1[4]]]
    bnd2 = ([BOUNDS_M1[0][2], BOUNDS_M1[0][3], BOUNDS_M1[0][4]], [BOUNDS_M1[1][2], BOUNDS_M1[1][3], BOUNDS_M1[1][4]])
    s2_ok = False; r2_s2 = None
    for gi, pg in enumerate(p0_opts2):
        try:
            popt2, pcov2 = curve_fit(stage2_model, V_large, J_large, p0=pg, bounds=bnd2, maxfev=100000, method='trf')
            Jp2 = stage2_model(V_large, *popt2)
            ssr2 = np.sum((J_large - Jp2)**2); sst2 = np.sum((J_large - np.mean(J_large))**2)
            r2_s2 = 1 - ssr2 / sst2 if sst2 > 0 else 0.0
            s2_ok = True; break
        except RuntimeError:
            pass
    if not s2_ok:
        gf = fit_model1(V, J, sample_name, T)
        if gf is None: return None
        Rsh_s2 = gf['popt'][2]; k_s2 = gf['popt'][3]; m_s2 = gf['popt'][4]
    else:
        Rsh_s2, k_s2, m_s2 = popt2
    p0_s3 = [J0_s1, A_s1, Rsh_s2, k_s2, m_s2]
    bounds_s3 = (
        [max(BOUNDS_M1[0][0], J0_s1*0.1), max(BOUNDS_M1[0][1], A_s1*0.7),
         max(BOUNDS_M1[0][2], Rsh_s2*0.1), max(BOUNDS_M1[0][3], k_s2*0.01),
         max(BOUNDS_M1[0][4], m_s2*0.7)],
        [min(BOUNDS_M1[1][0], J0_s1*10.0), min(BOUNDS_M1[1][1], A_s1*1.5),
         min(BOUNDS_M1[1][2], Rsh_s2*10.0), min(BOUNDS_M1[1][3], k_s2*100.0),
         min(BOUNDS_M1[1][4], m_s2*1.5)],
    )
    s3_ok = False
    try:
        popt, pcov = curve_fit(lambda V, *p: model1(V, *p, Vt=Vt), V, J, p0=p0_s3, bounds=bounds_s3, maxfev=50000, method='trf')
        perr = np.sqrt(np.diag(pcov))
        s3_ok = True
    except RuntimeError:
        pass
    if not s3_ok:
        popt = np.array(p0_s3); perr = np.full(5, np.nan)
    J_pred = model1(V, *popt, Vt=Vt)
    ss_res = np.sum((J - J_pred)**2)
    ss_tot = np.sum((J - np.mean(J))**2)
    r_squared = 1 - ss_res / ss_tot if ss_tot > 0 else 0.0
    return {'popt': popt, 'pcov': pcov if s3_ok else None, 'perr': perr,
            'r_squared': r_squared, 'Vt': Vt, 'model': 'eq1',
            'segmented': True, 'V_seg': V_seg,
            'stage1': {'J0': J0_s1, 'A': A_s1, 'r2': r2_s1},
            'stage2': {'Rsh': Rsh_s2, 'k': k_s2, 'm': m_s2, 'r2': r2_s2}}

def configure_plot_style(width_cm=8.5, double_column=False):
    width_inches = (17.5 if double_column else width_cm) / CM_PER_INCH
    cjk_fonts = ['STSong', 'SimSun', 'SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'KaiTi', 'FangSong']
    cjk_available = []
    from matplotlib.font_manager import findfont, FontProperties
    for f in cjk_fonts:
        try:
            fp = FontProperties(family=f)
            path = findfont(fp, fallback_to_default=False)
            if path: cjk_available.append(f)
        except Exception:
            pass
    sans_fonts = ['Arial', 'DejaVu Sans', 'Liberation Sans']
    matplotlib.rcParams.update({
        'font.family': 'sans-serif',
        'font.sans-serif': sans_fonts + cjk_available,
        'font.weight': 'bold',
        'font.size': 9, 'axes.labelsize': 9,
        'xtick.labelsize': 8, 'ytick.labelsize': 8, 'legend.fontsize': 7,
        'figure.dpi': 300, 'savefig.dpi': 300, 'savefig.bbox': 'tight', 'savefig.pad_inches': 0.05,
        'xtick.direction': 'in', 'ytick.direction': 'in',
        'xtick.major.size': 4, 'xtick.major.width': 0.8,
        'ytick.major.size': 4, 'ytick.major.width': 0.8,
        'xtick.minor.size': 2, 'xtick.minor.width': 0.6,
        'ytick.minor.size': 2, 'ytick.minor.width': 0.6,
        'xtick.top': True, 'ytick.right': True,
        'axes.linewidth': 0.8, 'lines.linewidth': 1.2,
        'lines.markersize': 4, 'lines.markeredgewidth': 0.5,
        'legend.frameon': True, 'legend.framealpha': 0.85,
        'legend.edgecolor': '#cccccc', 'legend.fancybox': False,
        'axes.grid': False,
        'mathtext.fontset': 'stix',
    })
    return width_inches

def plot_model_fit_single(V, J, fit, name, ax, model_type='eq48'):
    popt = fit['popt']; Vt = fit['Vt']
    Vs = np.linspace(V.min(), V.max(), 1000)
    if model_type == 'eq48':
        J_fit  = three_component_model(Vs, *popt, Vt=Vt)
        J_main = main_diode_current(Vs, popt[0], popt[1], Vt)
        J_ohm  = ohmic_current(Vs, popt[2])
        J_extra = tat_current(Vs, popt[3], popt[4], popt[5])
    else:
        J_fit  = model1(Vs, *popt, Vt=Vt)
        J_main = main_diode_current(Vs, popt[0], popt[1], Vt)
        J_ohm  = ohmic_current(Vs, popt[2])
        J_extra = non_ohmic_tunneling(Vs, popt[3], popt[4])
    ax.semilogy(V, np.abs(J), 'o', markersize=3, color=COLOR_PALETTE['data'], alpha=0.7,
                label='Data', zorder=5, markeredgecolor=COLOR_PALETTE['data'], markeredgewidth=0.3)
    ax.semilogy(Vs, np.abs(J_fit), '-', linewidth=1.5, alpha=0.7,
                color=COLOR_PALETTE['total_fit'], label='$J_{dark}$ fit', zorder=4)
    ax.semilogy(Vs, np.abs(J_main), '--', linewidth=1.0,
                color=COLOR_PALETTE['j_main'], label='$J_{main}$', zorder=3)
    ax.semilogy(Vs, np.abs(J_ohm), '-.', linewidth=1.0,
                color=COLOR_PALETTE['j_ohm'], label='$J_{Ohm}$', zorder=3)
    ax.semilogy(Vs, np.abs(J_extra), ':', linewidth=1.0,
                color=COLOR_PALETTE['j_non_tat'], label='$J_{TAT}$' if model_type == 'eq48' else '$J_{non}$', zorder=3)
    ax.set_xlabel('Voltage (V)', fontweight='bold', fontfamily='Arial')
    ax.set_ylabel('Current Density (A/cm$^{2}$)', fontweight='bold', fontfamily='Arial')
    ax.tick_params(top=False, right=False, which='both', labelsize=8)
    ax.legend(fontsize=6, loc='lower left', frameon=False, handlelength=1.5, ncol=1, prop={'weight': 'normal'})

def plot_model_comparison(Vc, fit_c, Vs, fit_s, ax, model_type='eq48'):
    Vt = fit_c['Vt']
    Vcs = np.linspace(Vc.min(), Vc.max(), 1000)
    Vss = np.linspace(Vs.min(), Vs.max(), 1000)
    Jm_c = main_diode_current(Vcs, fit_c['popt'][0], fit_c['popt'][1], Vt)
    Jm_s = main_diode_current(Vss, fit_s['popt'][0], fit_s['popt'][1], Vt)
    ax.semilogy(Vcs, np.abs(Jm_c), '-', linewidth=1.5, color=COLOR_PALETTE['data'],
                label='Control $J_{main}$')
    ax.semilogy(Vss, np.abs(Jm_s), '--', linewidth=1.5, color=COLOR_PALETTE['total_fit'],
                label='Sample $J_{main}$')
    ax.set_xlabel('Voltage (V)', fontweight='bold', fontfamily='Arial')
    ax.set_ylabel('$J_{main}$ (A/cm$^{2}$)', fontweight='bold', fontfamily='Arial')
    ax.tick_params(top=False, right=False, which='both', labelsize=8)
    ax.legend(fontsize=6, frameon=False, loc='lower left', handlelength=1.5, prop={'weight': 'normal'})

def plot_leakage_compare(Vc, fit_c, Vs, fit_s, ax, model_type='eq48'):
    pc, ps = fit_c['popt'], fit_s['popt']
    Vcs = np.linspace(Vc.min(), 0, 1000); Vss = np.linspace(Vs.min(), 0, 1000)
    if model_type == 'eq48':
        Lc = np.abs(ohmic_current(Vcs, pc[2]) + tat_current(Vcs, pc[3], pc[4], pc[5]))
        Ls = np.abs(ohmic_current(Vss, ps[2]) + tat_current(Vss, ps[3], ps[4], ps[5]))
        Oc = np.abs(ohmic_current(Vcs, pc[2]))
        Tc = np.abs(tat_current(Vcs, pc[3], pc[4], pc[5]))
    else:
        Lc = np.abs(ohmic_current(Vcs, pc[2]) + non_ohmic_tunneling(Vcs, pc[3], pc[4]))
        Ls = np.abs(ohmic_current(Vss, ps[2]) + non_ohmic_tunneling(Vss, ps[3], ps[4]))
        Oc = np.abs(ohmic_current(Vcs, pc[2]))
        Tc = np.abs(non_ohmic_tunneling(Vcs, pc[3], pc[4]))
    ax.semilogy(Vcs, Lc, '-', linewidth=1.5, color=COLOR_PALETTE['data'],
                label='Control ($J_{Ohm}$+$J_{extra}$)')
    ax.semilogy(Vss, Ls, '--', linewidth=1.5, color=COLOR_PALETTE['total_fit'],
                label='Sample ($J_{Ohm}$+$J_{extra}$)')
    ax.semilogy(Vcs, Oc, ':', linewidth=0.8, color=COLOR_PALETTE['j_ohm'], alpha=0.6, label='Ctrl $J_{Ohm}$')
    ax.semilogy(Vcs, Tc, ':', linewidth=0.5, color=COLOR_PALETTE['j_non_tat'], alpha=0.4, label='Ctrl $J_{extra}$')
    ax.set_xlabel('Voltage (V)', fontweight='bold', fontfamily='Arial')
    ax.set_ylabel('Leakage (A/cm$^{2}$)', fontweight='bold', fontfamily='Arial')
    ax.tick_params(top=False, right=False, which='both', labelsize=8)
    ax.legend(fontsize=6, frameon=False, loc='lower left', handlelength=1.5, prop={'weight': 'normal'})
    ax.set_xlim(-0.5, -0.2)

def save_figure_formats(fig, base_path, hd=False):
    saved = []
    for fmt, ext, dpi in [('svg','.svg',None), ('pdf','.pdf',None), ('png','.png',300)]:
        kw = {'format': fmt, 'bbox_inches': 'tight', 'facecolor': 'white', 'edgecolor': 'none'}
        if dpi: kw['dpi'] = dpi
        fig.savefig(base_path + ext, **kw)
        saved.append(base_path + ext)
    if hd:
        fig.savefig(base_path + '_600dpi.png', dpi=600, bbox_inches='tight', facecolor='white', edgecolor='none')
        saved.append(base_path + '_600dpi.png')
    return saved

def export_component_data(V, J, fit, device_name, output_dir, model_type='eq48'):
    popt = fit['popt']; Vt = fit['Vt']
    if model_type == 'eq48':
        J_fit = three_component_model(V, *popt, Vt=Vt)
        J_main = main_diode_current(V, popt[0], popt[1], Vt)
        J_ohm = ohmic_current(V, popt[2])
        J_extra = tat_current(V, popt[3], popt[4], popt[5])
        header = 'V(V)\tJ_data(A/cm2)\tJ_fit(A/cm2)\tJ_main(A/cm2)\tJ_Ohm(A/cm2)\tJ_TAT(A/cm2)'
    else:
        J_fit = model1(V, *popt, Vt=Vt)
        J_main = main_diode_current(V, popt[0], popt[1], Vt)
        J_ohm = ohmic_current(V, popt[2])
        J_extra = non_ohmic_tunneling(V, popt[3], popt[4])
        header = 'V(V)\tJ_data(A/cm2)\tJ_fit(A/cm2)\tJ_main(A/cm2)\tJ_Ohm(A/cm2)\tJ_non(A/cm2)'
    rows = np.column_stack([V, J, J_fit, J_main, J_ohm, J_extra])
    path = os.path.join(output_dir, f'{device_name}_component_fit.txt')
    np.savetxt(path, rows, header=header, delimiter='\t', fmt='%.6e', comments='')
    print(f'  Component data exported: {path}')
    return path

def export_parameters_table(fit_c, fit_s, output_dir, model_type='eq48'):
    pc, ps = fit_c['popt'], fit_s['popt']
    ec = fit_c.get('perr', np.full(len(pc), np.nan))
    es = fit_s.get('perr', np.full(len(ps), np.nan))
    if model_type == 'eq48':
        params = [
            ('J0', 'J0', 'A/cm2', pc[0], ec[0], ps[0], es[0]),
            ('A', 'A', '-', pc[1], ec[1], ps[1], es[1]),
            ('Rsh', 'Rsh', 'ohm.cm2', pc[2], ec[2], ps[2], es[2]),
            ('B', 'B', '-', pc[3], ec[3], ps[3], es[3]),
            ('C', 'C', '-', pc[4], ec[4], ps[4], es[4]),
            ('Vbi', 'Vbi', 'V', pc[5], ec[5], ps[5], es[5]),
        ]
        suffix = 'eq2'
    else:
        params = [
            ('J0', 'J0', 'A/cm2', pc[0], ec[0], ps[0], es[0]),
            ('A', 'A', '-', pc[1], ec[1], ps[1], es[1]),
            ('Rsh', 'Rsh', 'ohm.cm2', pc[2], ec[2], ps[2], es[2]),
            ('k', 'k', '-', pc[3], ec[3], ps[3], es[3]),
            ('m', 'm', '-', pc[4], ec[4], ps[4], es[4]),
        ]
        suffix = 'eq1'
    rows_display = []
    for label, sym, unit, cv, ce, sv, se in params:
        if unit == 'A/cm2' or sym == 'k':
            cv_str = f'{cv:.4e}'; sv_str = f'{sv:.4e}'
            ce_str = f'{ce:.4e}' if not np.isnan(ce) else '-'
            se_str = f'{se:.4e}' if not np.isnan(se) else '-'
        elif sym in ('Rsh',):
            cv_str = f'{cv:.2e}'; sv_str = f'{sv:.2e}'
            ce_str = f'{ce:.2e}' if not np.isnan(ce) else '-'
            se_str = f'{se:.2e}' if not np.isnan(se) else '-'
        else:
            cv_str = f'{cv:.4f}'; sv_str = f'{sv:.4f}'
            ce_str = f'{ce:.4f}' if not np.isnan(ce) else '-'
            se_str = f'{se:.4f}' if not np.isnan(se) else '-'
        rows_display.append([label, sym, unit, cv_str, ce_str, sv_str, se_str])
    rows_display.append(['-', '-', '-', '-', '-', '-', '-'])
    rows_display.append(['R2', 'R2', '-', f'{fit_c["r_squared"]:.6f}', '-', f'{fit_s["r_squared"]:.6f}', '-'])
    if fit_c.get('segmented'):
        s1c = fit_c.get('stage1', {}).get('r2'); s1s = fit_s.get('stage1', {}).get('r2')
        s2c = fit_c.get('stage2', {}).get('r2'); s2s = fit_s.get('stage2', {}).get('r2')
        if s1c is not None:
            rows_display.append(['Stage1 R2', 'R2_s1', '-', f'{s1c:.6f}', '-', f'{s1s:.6f}', '-'])
        if s2c is not None:
            rows_display.append(['Stage2 R2', 'R2_s2', '-', f'{s2c:.6f}', '-', f'{s2s:.6f}', '-'])
    df = pd.DataFrame(rows_display, columns=['Parameter', 'Symbol', 'Unit', 'Control', 'Ctrl +/-', 'Sample', 'Samp +/-'])
    print(f"\n{'='*70}\n  {model_type.upper()} Fitting Parameters\n{'='*70}")
    print(df.to_string(index=False))
    csv_p = os.path.join(output_dir, f'{suffix}_fitting_params.csv')
    df.to_csv(csv_p, index=False, encoding='utf-8-sig')
    print(f"  Params saved: {csv_p}")
    return df

def generate_combined_word_report(fit1_c, fit1_s, fit2_c, fit2_s, val1, val2, output_dir,
                                   control_label="Control PD", sample_label="Sample PD",
                                   T=300, area=None, V_seg=V_SEG_DEFAULT, ref_data=None):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("python-docx not installed."); return None
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    title = doc.add_heading('PbS CQD Photodetector Dark Current Component Fitting Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. Fitting Models', level=1)
    doc.add_heading('1.1 Model 1 - Eq.1', level=2)
    p = doc.add_paragraph(); p.add_run('J_dark = J0*[exp(qV/(A*kT))-1] + V/Rsh + k*V^m').font.name = 'Consolas'
    doc.add_heading('1.2 Model 2 - Eq.2 (TAT)', level=2)
    p2 = doc.add_paragraph(); p2.add_run('J_dark = J0*[exp(qV/(A*kT))-1] + V/Rsh + B*V*exp(-C/(Vbi-V))').font.name = 'Consolas'
    doc.add_heading('1.3 Segmented Fitting', level=2)
    doc.add_paragraph(f'Stage 1 (|V|<{V_seg}V): J0,A | Stage 2 (V<-{V_seg}V): remaining | Stage 3: global refinement')
    doc.add_heading('2. Model 1 (Eq.1) Results', level=1)
    _add_model_params_table(doc, fit1_c, fit1_s, 'eq1', control_label, sample_label)
    doc.add_heading('3. Model 2 (Eq.2) Results', level=1)
    _add_model_params_table(doc, fit2_c, fit2_s, 'eq48', control_label, sample_label)
    doc.add_heading('4. Model Comparison', level=1)
    r2_1c = fit1_c['r_squared']; r2_1s = fit1_s['r_squared']
    r2_2c = fit2_c['r_squared']; r2_2s = fit2_s['r_squared']
    t_comp = doc.add_table(rows=4, cols=5, style='Light Grid Accent 1')
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Metric', 'Model 1 Ctrl', 'Model 1 Samp', 'Model 2 Ctrl', 'Model 2 Samp']):
        t_comp.rows[0].cells[i].text = h
    t_comp.rows[1].cells[0].text = 'R2'; t_comp.rows[1].cells[1].text = f'{r2_1c:.6f}'
    t_comp.rows[1].cells[2].text = f'{r2_1s:.6f}'; t_comp.rows[1].cells[3].text = f'{r2_2c:.6f}'
    t_comp.rows[1].cells[4].text = f'{r2_2s:.6f}'
    t_comp.rows[2].cells[0].text = 'Complexity'; t_comp.rows[2].cells[1].text = '5 params'
    t_comp.rows[2].cells[2].text = '5 params'; t_comp.rows[2].cells[3].text = '6 params'
    t_comp.rows[2].cells[4].text = '6 params'
    t_comp.rows[3].cells[0].text = 'Physics'; t_comp.rows[3].cells[1].text = 'Phenomenological'
    t_comp.rows[3].cells[2].text = 'Phenomenological'; t_comp.rows[3].cells[3].text = 'Clear (TAT)'
    t_comp.rows[3].cells[4].text = 'Clear (TAT)'
    better_model = 'Model 2 (Eq.2)' if (r2_2c + r2_2s) > (r2_1c + r2_1s) else 'Model 1 (Eq.1)'
    doc.add_paragraph(f'Recommendation: {better_model} based on R2 and physical interpretability.')
    doc.add_heading('5. Device Optimization Mechanism', level=1)
    pc, ps = fit2_c['popt'], fit2_s['popt']
    doc.add_paragraph(f'J0: {pc[0]:.2e} -> {ps[0]:.2e} A/cm2 | A: {pc[1]:.2f} -> {ps[1]:.2f}')
    doc.add_paragraph(f'Rsh: {pc[2]:.2e} -> {ps[2]:.2e} ohm*cm2')
    doc.add_paragraph(f'B: {pc[3]:.2f} -> {ps[3]:.2f} | C: {pc[4]:.2f} -> {ps[4]:.2f}')
    report_path = os.path.join(output_dir, 'fitting_report.docx')
    doc.save(report_path)
    print(f"\n  Combined Word report saved: {report_path}")
    return report_path

def _add_model_params_table(doc, fit_c, fit_s, model_type, ctrl_label, samp_label):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    pc, ps = fit_c['popt'], fit_s['popt']
    ec = fit_c.get('perr', np.full(len(pc), np.nan))
    es = fit_s.get('perr', np.full(len(ps), np.nan))
    if model_type == 'eq48':
        param_data = [
            ('J0', 'A/cm2', 0, '.4e'), ('A', '-', 1, '.4f'),
            ('Rsh', 'ohm*cm2', 2, '.2e'), ('B', '-', 3, '.4f'),
            ('C', '-', 4, '.4f'), ('Vbi', 'V', 5, '.4f'),
        ]
    else:
        param_data = [
            ('J0', 'A/cm2', 0, '.4e'), ('A', '-', 1, '.4f'),
            ('Rsh', 'ohm*cm2', 2, '.2e'), ('k', '-', 3, '.4e'),
            ('m', '-', 4, '.4f'),
        ]
    n = len(param_data)
    t = doc.add_table(rows=n+2, cols=5, style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Param', 'Unit', ctrl_label, samp_label, 'Trend']):
        t.rows[0].cells[i].text = h
    for i, (pname, unit, idx, fmt_v) in enumerate(param_data):
        t.rows[i+1].cells[0].text = pname; t.rows[i+1].cells[1].text = unit
        cv = pc[idx]; sv = ps[idx]
        ce_v = ec[idx] if idx < len(ec) else np.nan
        se_v = es[idx] if idx < len(es) else np.nan
        t.rows[i+1].cells[2].text = f'{cv:{fmt_v}} +/- {ce_v:{fmt_v}}' if not np.isnan(ce_v) else f'{cv:{fmt_v}}'
        t.rows[i+1].cells[3].text = f'{sv:{fmt_v}} +/- {se_v:{fmt_v}}' if not np.isnan(se_v) else f'{sv:{fmt_v}}'
    t.rows[n+1].cells[0].text = 'R2'; t.rows[n+1].cells[1].text = '-'
    t.rows[n+1].cells[2].text = f'{fit_c["r_squared"]:.6f}'
    t.rows[n+1].cells[3].text = f'{fit_s["r_squared"]:.6f}'

def run_single_model(Vc, Jc, Vs, Js, output_dir, T, model_type, segmented, v_seg, validate, area, hd, fig_width_cm, double_column):
    print(f"\n# Fitting: {'Eq.2' if model_type == 'eq48' else 'Eq.1'}")
    fit_func = fit_dark_current if model_type == 'eq48' else fit_model1
    seg_func = segmented_fit_dark_current if model_type == 'eq48' else segmented_fit_model1
    if segmented:
        fit_c = seg_func(Vc, Jc, "Control", T, V_seg=v_seg)
        fit_s = seg_func(Vs, Js, "Sample", T, V_seg=v_seg)
    else:
        fit_c = fit_func(Vc, Jc, "Control", T)
        fit_s = fit_func(Vs, Js, "Sample", T)
    if fit_c is None or fit_s is None:
        print("Fitting failed."); return None
    configure_plot_style(width_cm=fig_width_cm, double_column=double_column)
    for device_name, V, J, fit in [('control', Vc, Jc, fit_c), ('sample', Vs, Js, fit_s)]:
        fig, ax = plt.subplots(figsize=(fig_width_cm/CM_PER_INCH, fig_width_cm/CM_PER_INCH*0.75))
        plot_model_fit_single(V, J, fit, '', ax, model_type)
        plt.tight_layout()
        save_figure_formats(fig, os.path.join(output_dir, f'{device_name}_fitting'), hd=hd)
        plt.close(fig)
    for comp_name, comp_fn in [('j_main_comparison', plot_model_comparison), ('leakage_comparison', plot_leakage_compare)]:
        fig, ax = plt.subplots(figsize=(fig_width_cm/CM_PER_INCH, fig_width_cm/CM_PER_INCH*0.75))
        comp_fn(Vc, fit_c, Vs, fit_s, ax, model_type)
        plt.tight_layout()
        save_figure_formats(fig, os.path.join(output_dir, comp_name), hd=hd)
        plt.close(fig)
    export_component_data(Vc, Jc, fit_c, 'control', output_dir, model_type)
    export_component_data(Vs, Js, fit_s, 'sample', output_dir, model_type)
    export_parameters_table(fit_c, fit_s, output_dir, model_type)
    return {'control': fit_c, 'sample': fit_s}

def run_dark_current_fitting(control_file, sample_file, output_dir=None, T=300, area=None,
                              vmin=None, vmax=None, auto_dark=True, sweep_c=None, sweep_s=None,
                              segmented=True, v_seg=V_SEG_DEFAULT, validate=True, generate_report=True,
                              report_format='docx', model='both', max_points=None,
                              hd=False, fig_width_cm=8.5, double_column=False):
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(control_file)), 'output')
    os.makedirs(output_dir, exist_ok=True)
    models_to_run = ['eq1', 'eq48'] if model == 'both' else [model]
    print(f"\nPbS CQD Dark Current Fitting | Models: {models_to_run}")
    Vc, Jc = load_data(control_file, area=area, auto_dark=auto_dark, manual_sweep=sweep_c, max_points=max_points)
    Vs, Js = load_data(sample_file, area=area, auto_dark=auto_dark, manual_sweep=sweep_s, max_points=max_points)
    if vmin is not None or vmax is not None:
        lo = vmin if vmin is not None else Vc.min()
        hi = vmax if vmax is not None else Vc.max()
        mc = (Vc >= lo) & (Vc <= hi); Vc, Jc = Vc[mc], Jc[mc]
        ms = (Vs >= lo) & (Vs <= hi); Vs, Js = Vs[ms], Js[ms]
    results = {}
    for m in models_to_run:
        model_dir = os.path.join(output_dir, f'model_{"eq1" if m == "eq1" else "eq2"}' if model == 'both' else output_dir)
        if model == 'both':
            os.makedirs(model_dir, exist_ok=True)
            od = model_dir
        else:
            od = output_dir
        r = run_single_model(Vc.copy(), Jc.copy(), Vs.copy(), Js.copy(), od, T, m, segmented, v_seg, validate, area, hd, fig_width_cm, double_column)
        if r: results[m] = r
    if not results:
        print("All fits failed."); return None
    if generate_report and len(results) >= 2:
        r1 = results.get('eq1', {}); r2 = results.get('eq48', {})
        if r1 and r2:
            generate_combined_word_report(r1['control'], r1['sample'], r2['control'], r2['sample'],
                {'all_valid': True, 'violation_count': 0, 'total_params': 4},
                {'all_valid': True, 'violation_count': 0, 'total_params': 5},
                output_dir, T=T, area=area, V_seg=v_seg, ref_data=REFERENCE_DATA)
    print("\nDone!")
    return {'results': results, 'output_dir': output_dir}

if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser(description='PbS CQD Dark Current Component Fitting')
    p.add_argument('control', help='Control PD data file')
    p.add_argument('sample', help='Sample PD data file')
    p.add_argument('-o', '--output', default=None); p.add_argument('-T', '--temperature', type=float, default=T_DEFAULT)
    p.add_argument('-a', '--area', type=float, default=None)
    p.add_argument('--vmin', type=float, default=-0.5); p.add_argument('--vmax', type=float, default=0.2)
    p.add_argument('--no-auto-dark', action='store_true')
    p.add_argument('--model', choices=['both', 'eq1', 'eq48'], default='both')
    p.add_argument('--points', type=int, default=None)
    p.add_argument('--no-segmented', action='store_true')
    p.add_argument('--vseg', type=float, default=V_SEG_DEFAULT)
    p.add_argument('--no-validate', action='store_true')
    p.add_argument('--no-report', action='store_true')
    p.add_argument('--fig-width-cm', type=float, default=8.5)
    p.add_argument('--fig-double-col', action='store_true')
    p.add_argument('--hd', action='store_true')
    args = p.parse_args()
    r = run_dark_current_fitting(args.control, args.sample, args.output,
        args.temperature, args.area, args.vmin, args.vmax,
        auto_dark=not args.no_auto_dark,
        sweep_c=args.sweep_control if hasattr(args, 'sweep_control') else None,
        sweep_s=args.sweep_sample if hasattr(args, 'sweep_sample') else None,
        segmented=not args.no_segmented, v_seg=args.vseg,
        validate=not args.no_validate,
        generate_report=not args.no_report,
        model=args.model, max_points=args.points,
        hd=args.hd, fig_width_cm=args.fig_width_cm, double_column=args.fig_double_col)
    if r: print(f"\nOutput: {r['output_dir']}")